import os
import json
import time
from . import config
from docx import Document
from docx.shared import Inches

def compile_final_report():
    """
    Reads all individual JSON analysis files and compiles them into a single,
    professional Word document (.docx) with a formatted table.
    """
    print("\n[Compiler] Starting final report compilation...")
    
    final_reports_dir = config.PROJECT_ROOT / "data" / "final_reports"
    final_reports_dir.mkdir(exist_ok=True)
    
    report_path = final_reports_dir / "Consolidated_Security_Analysis_Report.docx"
    json_source_dir = config.OUTPUT_REPORTS_DIR

    try:
        json_files = sorted([f for f in os.listdir(json_source_dir) if f.endswith('_analysis.json')])
        if not json_files:
            print("[Compiler] No JSON analysis files found to compile.")
            return
    except FileNotFoundError:
        print(f"[Compiler] ERROR: Source directory not found: {json_source_dir}")
        return

    # --- NEW: Word Document and Table Creation ---
    document = Document()
    document.add_heading('Consolidated Security Analysis Report', level=1)
    document.add_paragraph(f"This report summarizes the automated analysis of {len(json_files)} processed files.")
    
    # Define table with 1 row for headers
    table = document.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    # Populate Header Row
    hdr_cells = table.rows[0].cells
    headers = ["File Name", "File Type", "File Description", "Key Security Findings"]
    for i, header_text in enumerate(headers):
        hdr_cells[i].text = header_text
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True

    # Populate Data Rows
    for file_name in json_files:
        try:
            with open(os.path.join(json_source_dir, file_name), 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            row_cells = table.add_row().cells
            row_cells[0].text = data.get("source_file_name", "N/A")
            row_cells[1].text = data.get("source_file_type", "N/A")
            row_cells[2].text = data.get("document_summary", "No summary provided.")
            
            # Add findings as a bulleted list within the cell
            findings_list = data.get("key_security_findings", ["No findings reported."])
            # Clear the default paragraph in the cell before adding new ones
            findings_cell = row_cells[3]
            findings_cell.text = "" # Clear any default text
            for item in findings_list:
                # Add each finding as a new paragraph with a bullet point style
                p = findings_cell.add_paragraph(item)
                p.style = 'List Bullet'

        except (json.JSONDecodeError, KeyError) as e:
            print(f"[Compiler] WARNING: Could not process {file_name}. Error: {e}")

    # --- Save the Word Document with Retry Logic ---
    max_retries = 3
    for attempt in range(max_retries):
        try:
            document.save(report_path)
            print(f"✅ [Compiler] Successfully created Word report: {report_path.name}")
            return
        except PermissionError:
            if attempt < max_retries - 1:
                print(f"   -> ❌ [Compiler] ERROR: Permission denied for {report_path.name}. The file is likely open.")
                print(f"      Please close the file. Retrying in 10 seconds... (Attempt {attempt + 2}/{max_retries})")
                time.sleep(10)
            else:
                print(f"   -> ❌ [Compiler] CRITICAL: Failed to save report after {max_retries} attempts. Please ensure the file is closed.")
                return